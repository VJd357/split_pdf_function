"""
PDF Document Splitter

This module provides functionality to split PDF documents into sections based on headings
and save them in various formats (PDF, DOCX, TXT).

Classes:
    DocumentSection: Represents a section of a document with content and metadata
    DocumentSplitter: Main class for splitting PDFs into sections based on headings
"""

import PyPDF2
import pdfplumber
from docx import Document
import re
import os
from pathlib import Path
from fuzzywuzzy import fuzz
import logging
from typing import List, Dict, Tuple, Optional
from dataclasses import dataclass
import json

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


@dataclass
class DocumentSection:
    """
    Represents a section of a document with its content and metadata.

    Attributes:
        heading (str): Title/heading of the section
        pages (List[int]): List of page numbers in this section
        content (List[str]): List of text content from the pages
        confidence (float): Confidence score of heading match (0-100)
        start_page (int): First page number of section
        end_page (int): Last page number of section
        pdf_path (str): Path to source PDF file
        is_fuzzy_match (bool): Whether heading was matched using fuzzy matching
    """
    heading: str
    pages: List[int]
    content: List[str]
    confidence: float
    start_page: int
    end_page: int
    pdf_path: str = ""
    is_fuzzy_match: bool = False


class DocumentSplitter:
    """
    Splits PDF documents into sections based on heading detection.

    The splitter uses exact and fuzzy matching to identify document sections
    and can output them in PDF, DOCX or TXT format.

    Args:
        min_confidence (float): Minimum confidence threshold for heading matches (0-100)
        fuzzy_threshold (float): Minimum threshold for fuzzy matching (0-100)
    """

    def __init__(self, min_confidence: float = 80, fuzzy_threshold: float = 80):
        self.min_confidence = min_confidence
        self.fuzzy_threshold = fuzzy_threshold
        
        # Suppress pdfplumber warnings
        logging.getLogger('pdfminer').setLevel(logging.ERROR)
        
        # Define document groups for fuzzy matching
        self.document_groups = {
            "Self Declaration Cum Undertaking Certificate": [
                "Self Declaration Cum Undertaking Certificate",
                "Declaration Cum Undertaking Certificate", 
                "Self Declaration Certificate",
                "Declaration Certificate"
            ],
            "Certificate of Origin": [
                "Certificate of Origin",
                "Origin Certificate"
            ],
            "Certificate of Chemical Analysis Report": [
                "Certificate of Chemical Analysis Report",
                "Chemical Analysis Report",
                "Chemical Analysis Certificate"
            ],
            "Form 6": ["Form 6"],
            "Form 9": ["Form 9"],
            "Bill of Exchange": ["Bill of Exchange"],
            "Commercial Invoice": ["Commercial Invoice"],
            "Packing List": ["Packing List"],
            "Pre-Shipment Inspection Certificate": ["Pre-Shipment Inspection Certificate"],
            "Insurance Policy": ["Insurance Policy"],
            "Transboundary Movement Document": ["Transboundary Movement Document"],
            "Bill of Lading": ["Bill of Lading"]
        }
        
        # Create heading patterns for better matching
        self.heading_patterns = {}
        for main_heading, variations in self.document_groups.items():
            patterns = []
            for variation in variations:
                patterns.extend([
                    re.escape(variation),
                    variation.replace(" ", ".*"),
                    variation.replace(" ", "\\s*"),
                    variation.upper(),
                    variation.lower()
                ])
            self.heading_patterns[main_heading] = patterns

    def normalize_text(self, text: str) -> str:
        """
        Normalize text for consistent comparison.

        Args:
            text (str): Text to normalize

        Returns:
            str: Normalized text (uppercase with normalized whitespace)
        """
        return re.sub(r'\s+', ' ', text.strip().upper())

    def find_heading(self, text: str) -> Tuple[Optional[str], float, bool]:
        """
        Find the best matching heading for the given text.

        Args:
            text (str): Text to search for headings in

        Returns:
            Tuple containing:
                - str or None: Matched heading or None if no match
                - float: Confidence score of the match (0-100)
                - bool: Whether fuzzy matching was used
        """
        normalized_text = self.normalize_text(text)
        best_match = None
        best_confidence = 0.0
        is_fuzzy = False

        for main_heading, patterns in self.heading_patterns.items():
            # Try exact match first
            if normalized_text == self.normalize_text(main_heading):
                return main_heading, 100.0, False

            # Try pattern matching
            for pattern in patterns:
                if re.search(rf'\b{pattern}\b', text, re.IGNORECASE):
                    return main_heading, 95.0, False

            # Try fuzzy matching
            for variation in self.document_groups[main_heading]:
                confidence = fuzz.partial_ratio(normalized_text, self.normalize_text(variation))
                if confidence > best_confidence:
                    best_confidence = confidence
                    best_match = main_heading
                    is_fuzzy = True

        return best_match, best_confidence, is_fuzzy

    def extract_page_content(self, page) -> List[str]:
        """
        Extract text and table content from a PDF page.

        Args:
            page: pdfplumber page object

        Returns:
            List[str]: List of text lines from the page
        """
        content = []
        
        # Extract text
        text = page.extract_text()
        if text:
            content.extend(text.split('\n'))
        
        # Extract tables
        tables = page.extract_tables()
        for table in tables:
            for row in table:
                content.append(' | '.join(str(cell) if cell else '' for cell in row))
        
        return [line.strip() for line in content if line.strip()]

    def process_pdf(self, pdf_path: str) -> List[DocumentSection]:
        """
        Process PDF and split into sections based on headings.

        Uses a two-pass approach:
        1. Find sections with exact heading matches
        2. Process remaining pages with fuzzy matching
        Creates an "Others" section for unmatched pages.

        Args:
            pdf_path (str): Path to PDF file to process

        Returns:
            List[DocumentSection]: List of document sections found
        """
        sections = []
        current_section = None
        processed_pages = set()
        
        # First pass: Process pages with exact matches
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                content = self.extract_page_content(page)
                
                # Check first few lines for heading
                heading_found = False
                for line in content[:5]:
                    heading, confidence, is_fuzzy = self.find_heading(line)
                    if heading and confidence >= self.min_confidence and not is_fuzzy:
                        if current_section:
                            current_section.end_page = page_num
                            sections.append(current_section)
                        
                        current_section = DocumentSection(
                            heading=heading,
                            pages=[page_num],
                            content=content,
                            confidence=confidence,
                            start_page=page_num,
                            end_page=page_num,
                            pdf_path=pdf_path,
                            is_fuzzy_match=False
                        )
                        processed_pages.add(page_num)
                        heading_found = True
                        break
                
                if not heading_found and current_section:
                    current_section.pages.append(page_num)
                    current_section.content.extend(content)
                    current_section.end_page = page_num
                    processed_pages.add(page_num)
        
        if current_section:
            sections.append(current_section)
        
        # Second pass: Process remaining pages with fuzzy matching
        with pdfplumber.open(pdf_path) as pdf:
            current_section = None
            for page_num, page in enumerate(pdf.pages):
                if page_num in processed_pages:
                    continue
                    
                content = self.extract_page_content(page)
                best_heading = None
                best_confidence = 0.0
                
                # Check all lines for fuzzy matches
                for line in content:
                    heading, confidence, is_fuzzy = self.find_heading(line)
                    if heading and is_fuzzy and confidence > best_confidence:
                        best_heading = heading
                        best_confidence = confidence
                
                if best_heading and best_confidence >= self.fuzzy_threshold:
                    if current_section:
                        current_section.end_page = page_num - 1
                        sections.append(current_section)
                    
                    current_section = DocumentSection(
                        heading=best_heading,
                        pages=[page_num],
                        content=content,
                        confidence=best_confidence,
                        start_page=page_num,
                        end_page=page_num,
                        pdf_path=pdf_path,
                        is_fuzzy_match=True
                    )
                    processed_pages.add(page_num)
                elif current_section:
                    current_section.pages.append(page_num)
                    current_section.content.extend(content)
                    current_section.end_page = page_num
                    processed_pages.add(page_num)
        
        if current_section:
            sections.append(current_section)
        
        # Create "Others" section for remaining pages
        if processed_pages:
            other_pages = sorted(set(range(len(pdf.pages))) - processed_pages)
            if other_pages:
                other_content = []
                for page_num in other_pages:
                    content = self.extract_page_content(pdf.pages[page_num])
                    other_content.extend(content)
                
                others_section = DocumentSection(
                    heading="Others",
                    pages=other_pages,
                    content=other_content,
                    confidence=100.0,
                    start_page=other_pages[0],
                    end_page=other_pages[-1],
                    pdf_path=pdf_path,
                    is_fuzzy_match=False
                )
                sections.append(others_section)
        
        return sections

    def save_sections(
        self, 
        sections: List[DocumentSection], 
        output_dir: str, 
        base_filename: str, 
        format: str = "pdf"
    ) -> None:
        """
        Save document sections as separate files.

        Args:
            sections (List[DocumentSection]): Sections to save
            output_dir (str): Directory to save files in
            base_filename (str): Base name for output files
            format (str): Output format - "pdf", "docx" or "txt"
        """
        os.makedirs(output_dir, exist_ok=True)
        
        for section in sections:
            try:
                safe_heading = re.sub(r'[^\w\s-]', '', section.heading).replace(' ', '_')
                output_path = os.path.join(
                    output_dir, 
                    f"{base_filename}_{safe_heading}.{format}"
                )
                
                if format == "pdf":
                    self._save_as_pdf(section, output_path)
                elif format == "docx":
                    self._save_as_docx(section, output_path)
                else:  # txt
                    self._save_as_txt(section, output_path)
                
                match_type = "fuzzy" if section.is_fuzzy_match else "exact"
                logger.info(
                    f"Saved section '{section.heading}' ({match_type} match) to {output_path}"
                )
            except Exception as e:
                logger.error(f"Failed to save section '{section.heading}': {e}")
                continue

    def _save_as_pdf(self, section: DocumentSection, output_path: str) -> None:
        """
        Save section as PDF.

        Args:
            section (DocumentSection): Section to save
            output_path (str): Output file path
        """
        try:
            with open(section.pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                writer = PyPDF2.PdfWriter()
                
                # Add pages to the new PDF
                for page_num in section.pages:
                    if page_num < len(reader.pages):
                        writer.add_page(reader.pages[page_num])
                
                # Write the output PDF
                with open(output_path, 'wb') as output_file:
                    writer.write(output_file)
                    
            logger.info(f"Successfully saved PDF section: {section.heading}")
        except Exception as e:
            logger.error(f"Error saving PDF section: {e}")
            raise

    def _save_as_docx(self, section: DocumentSection, output_path: str) -> None:
        """
        Save section as DOCX.

        Args:
            section (DocumentSection): Section to save
            output_path (str): Output file path
        """
        try:
            doc = Document()
            doc.add_heading(section.heading, level=1)
            for line in section.content:
                doc.add_paragraph(line)
            doc.save(output_path)
        except Exception as e:
            logger.error(f"Error saving DOCX section: {e}")

    def _save_as_txt(self, section: DocumentSection, output_path: str) -> None:
        """
        Save section as TXT.

        Args:
            section (DocumentSection): Section to save
            output_path (str): Output file path
        """
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(f"{section.heading}\n\n")
                f.write('\n'.join(section.content))
        except Exception as e:
            logger.error(f"Error saving TXT section: {e}")


def main():
    """Main entry point for the PDF splitter tool."""
    # Get input file path from user
    input_file = input("Please enter the path to your PDF file: ").strip()
    output_directory = Path(input_file).stem
    output_format = "pdf"  # Options: "txt", "pdf", "docx"
    
    try:
        splitter = DocumentSplitter(min_confidence=80, fuzzy_threshold=80)
        sections = splitter.process_pdf(input_file)
        splitter.save_sections(sections, output_directory, Path(input_file).stem, output_format)
    except Exception as e:
        logger.error(f"Error processing document: {e}")


if __name__ == "__main__":
    main()