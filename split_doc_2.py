import PyPDF2
import pdfplumber
from docx import Document
import pytesseract
from PIL import Image
import io
import re
import os
from pathlib import Path
from fuzzywuzzy import fuzz
import logging

pytesseract.pytesseract.tesseract_cmd = r'C:\Users\akatre\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'


# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# List of headings to search for
HEADINGS = [
    "Commercial Invoice",
    "Packing List",
    "Certificate of Origin",
    "Self Declaration Cum Undertaking Certificate",
    "Pre-Shipment Inspection Certificate",
    "Insurance Policy",
    "Transboundary Movement Document",
    "Bill of Lading",
    "Certificate of Chemical Analysis Report"
]

# Unified configuration
CONFIG = {
    "heading_patterns": {
        h: [re.escape(h), h.replace(" ", ".*"), h.replace(" ", "\\s*")] for h in HEADINGS
    },
    "min_similarity": 80,
    # "ocr_required": False  # Set to True for scanned PDFs
}

def normalize_heading(text):
    """Normalize text for heading comparison."""
    return re.sub(r'\s+', ' ', text.strip().upper())

def extract_text_from_pdf(pdf_path, config):
    """Extract text from PDF with layout awareness and optional OCR."""
    pages = []
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            # Try text extraction first
            text = page.extract_text()
            if not text or config.get("ocr_required"):
                logger.info(f"Using OCR for page {page_num + 1}")
                try:
                    image = page.to_image().original
                    text = pytesseract.image_to_string(image, lang='eng')
                except Exception as e:
                    logger.warning(f"OCR failed on page {page_num + 1}: {e}")
                    text = ""
            # Extract words with layout info
            words = page.extract_words()
            page_content = {
                "text": text.split('\n') if text else [],
                "words": [(w.get('text', ''), w.get('size', 10), w.get('x0', 0), w.get('y0', 0)) for w in words],
                "page_num": page_num
            }
            pages.append(page_content)
    return pages

def extract_text_from_docx(docx_path):
    """Extract text from DOCX with style information."""
    doc = Document(docx_path)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            style = para.style.name if para.style else "Normal"
            paragraphs.append({"text": para.text, "style": style})
    # Handle tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append({"text": cell.text, "style": "Table"})
    return paragraphs

def is_heading(text, heading_patterns, min_similarity, words=None):
    """Check if text matches a heading using fuzzy matching and layout cues."""
    normalized_text = normalize_heading(text)
    for heading, patterns in heading_patterns.items():
        # Fuzzy matching
        if fuzz.partial_ratio(normalized_text, normalize_heading(heading)) >= min_similarity:
            return heading
        # Regex matching
        for pattern in patterns:
            if re.match(rf'^{pattern}\b', text, re.IGNORECASE):
                return heading
        # Layout-based check (for PDFs)
        if words:
            try:
                # Check if any word in the text has a larger font size
                font_sizes = [w[1] for w in words if w[0] and normalize_heading(w[0]) in normalized_text]
                if font_sizes and max(font_sizes, default=10) > 12:  # Arbitrary threshold
                    return heading
            except Exception as e:
                logger.warning(f"Error in layout-based heading check: {e}")
    return None

def find_headings_in_content(content, heading_patterns, min_similarity):
    """Find headings in content and return their positions."""
    heading_positions = []
    for i, item in enumerate(content):
        text = item["text"] if isinstance(item, dict) else item
        if isinstance(text, list):
            text = ' '.join(text)
        words = item.get("words") if isinstance(item, dict) else None
        matched_heading = is_heading(text, heading_patterns, min_similarity, words)
        if matched_heading:
            page_num = item.get("page_num", 0) if isinstance(item, dict) else 0
            heading_positions.append((i, page_num, matched_heading))
    return heading_positions

def split_content_by_headings(content, heading_positions):
    """Split content into sections based on heading positions."""
    sections = {}
    for i, (start_idx, start_page, heading) in enumerate(heading_positions):
        end_idx = heading_positions[i + 1][0] if i + 1 < len(heading_positions) else len(content)
        end_page = heading_positions[i + 1][1] if i + 1 < len(heading_positions) else content[-1].get("page_num", 0) + 1
        # Collect content for this section
        section_content = []
        for j in range(start_idx, end_idx):
            item = content[j]
            text = item["text"] if isinstance(item, dict) else item
            if isinstance(text, list):
                text = ' '.join(text)
            section_content.append(text)
        sections[heading] = {
            "content": section_content,
            "page_range": (start_page, end_page)
        }
    return sections

def save_sections_as_pdf(sections, input_pdf_path, output_dir, base_filename):
    """Save sections as separate PDF files."""
    os.makedirs(output_dir, exist_ok=True)
    try:
        input_pdf = PyPDF2.PdfReader(input_pdf_path)
    except Exception as e:
        logger.error(f"Failed to read PDF: {e}")
        return
    
    for heading, data in sections.items():
        safe_heading = re.sub(r'[^\w\s-]', '', heading).replace(' ', '_')
        output_path = os.path.join(output_dir, f"{base_filename}_{safe_heading}.pdf")
        output_pdf = PyPDF2.PdfWriter()
        
        start_page, end_page = data["page_range"]
        for page_num in range(start_page, min(end_page, len(input_pdf.pages))):
            output_pdf.add_page(input_pdf.pages[page_num])
        
        with open(output_path, 'wb') as f:
            output_pdf.write(f)
        logger.info(f"Saved section '{heading}' to {output_path}")

def save_sections_as_docx(sections, output_dir, base_filename):
    """Save sections as separate DOCX files."""
    os.makedirs(output_dir, exist_ok=True)
    
    for heading, data in sections.items():
        safe_heading = re.sub(r'[^\w\s-]', '', heading).replace(' ', '_')
        output_path = os.path.join(output_dir, f"{base_filename}_{safe_heading}.docx")
        doc = Document()
        doc.add_heading(heading, level=1)
        for line in data["content"]:
            doc.add_paragraph(line)
        doc.save(output_path)
        logger.info(f"Saved section '{heading}' to {output_path}")

def process_document(file_path, output_dir, output_format="txt"):
    """Process a PDF or DOCX file and split by headings."""
    file_path = Path(file_path)
    base_filename = file_path.stem
    output_dir = Path(output_dir)

    # Extract content
    try:
        if file_path.suffix.lower() == '.pdf':
            content = extract_text_from_pdf(file_path, CONFIG)
        elif file_path.suffix.lower() == '.docx':
            content = extract_text_from_docx(file_path)
        else:
            raise ValueError("Unsupported file format. Use PDF or DOCX.")
    except Exception as e:
        logger.error(f"Failed to extract content: {e}")
        return

    # Find headings
    heading_positions = find_headings_in_content(content, CONFIG["heading_patterns"], CONFIG["min_similarity"])
    if not heading_positions:
        logger.warning("No headings found in the document.")
        return

    # Split content by headings
    sections = split_content_by_headings(content, heading_positions)

    # Save sections
    try:
        if output_format == "pdf" and file_path.suffix.lower() == '.pdf':
            save_sections_as_pdf(sections, file_path, output_dir, base_filename)
        elif output_format == "docx":
            save_sections_as_docx(sections, output_dir, base_filename)
        else:
            os.makedirs(output_dir, exist_ok=True)
            for heading, data in sections.items():
                safe_heading = re.sub(r'[^\w\s-]', '', heading).replace(' ', '_')
                output_path = os.path.join(output_dir, f"{base_filename}_{safe_heading}.txt")
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write('\n'.join(data["content"]))
                logger.info(f"Saved section '{heading}' to {output_path}")
    except Exception as e:
        logger.error(f"Failed to save sections: {e}")

def main():
    # Example usage
    input_file = r"C:\Users\akatre\Desktop\Gen AI Agents\Z Nice Import Agent\imports model\Imports billing Doc\CANUSA HERSHMAN RECYCLING COMPANY.pdf"  # Replace with your file path
    output_directory = "split_sections_advanced_5"
    output_format = "pdf"  # Options: "txt", "pdf", "docx"
    
    try:
        process_document(input_file, output_directory, output_format)
    except Exception as e:
        logger.error(f"Error processing document: {e}")

if __name__ == "__main__":
    main()